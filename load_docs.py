"""DOCX 문서를 읽어서 ChromaDB에 벡터로 저장하는 스크립트"""
import os
import docx
import chromadb
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()
client = OpenAI()

DOCS_DIR = os.path.join(os.path.dirname(__file__), "..", "농도계")
CHROMA_DIR = os.path.join(os.path.dirname(__file__), "chroma_db")


def extract_text_from_docx(filepath):
    """DOCX 파일에서 텍스트 추출 (단락 + 테이블)"""
    doc = docx.Document(filepath)
    texts = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            texts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                texts.append(" | ".join(cells))

    return "\n".join(texts)


def chunk_text(text, chunk_size=500, overlap=50):
    """텍스트를 일정 크기로 분할"""
    lines = text.split("\n")
    chunks = []
    current_chunk = []
    current_len = 0

    for line in lines:
        line_len = len(line)
        if current_len + line_len > chunk_size and current_chunk:
            chunks.append("\n".join(current_chunk))
            # overlap: 마지막 몇 줄 유지
            overlap_lines = []
            overlap_len = 0
            for prev_line in reversed(current_chunk):
                if overlap_len + len(prev_line) > overlap:
                    break
                overlap_lines.insert(0, prev_line)
                overlap_len += len(prev_line)
            current_chunk = overlap_lines
            current_len = overlap_len

        current_chunk.append(line)
        current_len += line_len

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks


def get_embedding(text):
    """OpenAI 임베딩 생성"""
    response = client.embeddings.create(
        model="text-embedding-3-small",
        input=text
    )
    return response.data[0].embedding


def extract_text_from_md(filepath):
    """MD 파일에서 텍스트 추출"""
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()


def load_files(base_dir, collection):
    """DOCX + MD 파일을 찾아서 벡터DB에 저장"""
    files = []

    # DOCX 파일 (상위 폴더)
    for f in os.listdir(base_dir):
        if f.lower().endswith(".docx"):
            files.append(("docx", os.path.join(base_dir, f)))

    # MD 파일 (docs 폴더)
    docs_dir = os.path.join(os.path.dirname(__file__), "docs")
    if os.path.exists(docs_dir):
        for f in os.listdir(docs_dir):
            if f.lower().endswith(".md"):
                files.append(("md", os.path.join(docs_dir, f)))

    if not files:
        print("파일을 찾을 수 없습니다.")
        return 0

    print(f"발견된 문서: {len(files)}개")
    for ftype, fpath in files:
        print(f"  - [{ftype}] {os.path.basename(fpath)}")

    total_chunks = 0
    for ftype, filepath in files:
        filename = os.path.basename(filepath)
        print(f"\n처리 중: {filename}")

        if ftype == "docx":
            text = extract_text_from_docx(filepath)
        else:
            text = extract_text_from_md(filepath)

        print(f"  텍스트 길이: {len(text)}자")

        chunks = chunk_text(text)
        print(f"  청크 수: {len(chunks)}개")

        for i, chunk in enumerate(chunks):
            chunk_id = f"{filename}_{i}"
            embedding = get_embedding(chunk)

            collection.add(
                ids=[chunk_id],
                embeddings=[embedding],
                documents=[chunk],
                metadatas=[{"source": filename, "chunk_index": i}]
            )

            if (i + 1) % 10 == 0:
                print(f"  {i + 1}/{len(chunks)} 임베딩 완료...")

        total_chunks += len(chunks)
        print(f"  완료!")

    return total_chunks


def main():
    # ChromaDB 초기화
    chroma_client = chromadb.PersistentClient(path=CHROMA_DIR)

    # 기존 컬렉션 삭제 후 재생성
    try:
        chroma_client.delete_collection("wess_density")
    except Exception:
        pass
    collection = chroma_client.create_collection(
        name="wess_density",
        metadata={"hnsw:space": "cosine"}
    )

    total = load_files(DOCS_DIR, collection)
    print(f"\n농도계 완료! 총 {total}개 청크 저장됨")

    # 계면계 컬렉션
    INTERFACE_DIR = os.path.join(os.path.dirname(__file__), "..", "계면계")
    if os.path.exists(INTERFACE_DIR):
        try:
            chroma_client.delete_collection("wess_interface")
        except Exception:
            pass
        interface_col = chroma_client.create_collection(
            name="wess_interface",
            metadata={"hnsw:space": "cosine"}
        )

        # 계면계 DOCX + MD 파일 로드
        ifiles = []
        for f in os.listdir(INTERFACE_DIR):
            if f.lower().endswith(".docx"):
                ifiles.append(("docx", os.path.join(INTERFACE_DIR, f)))

        # 계면계 MD 파일 (docs_interface 폴더)
        idocs_dir = os.path.join(os.path.dirname(__file__), "docs_interface")
        if os.path.exists(idocs_dir):
            for f in os.listdir(idocs_dir):
                if f.lower().endswith(".md"):
                    ifiles.append(("md", os.path.join(idocs_dir, f)))

        if ifiles:
            print(f"\n=== 계면계 문서: {len(ifiles)}개 ===")
            for ftype, fpath in ifiles:
                print(f"  - {os.path.basename(fpath)}")

            itotal = 0
            for ftype, filepath in ifiles:
                filename = os.path.basename(filepath)
                print(f"\n처리 중: {filename}")
                if ftype == "docx":
                    text = extract_text_from_docx(filepath)
                else:
                    text = extract_text_from_md(filepath)
                print(f"  텍스트 길이: {len(text)}자")
                chunks = chunk_text(text)
                print(f"  청크 수: {len(chunks)}개")
                for i, chunk in enumerate(chunks):
                    chunk_id = f"{filename}_{i}"
                    embedding = get_embedding(chunk)
                    interface_col.add(
                        ids=[chunk_id],
                        embeddings=[embedding],
                        documents=[chunk],
                        metadatas=[{"source": filename, "chunk_index": i}]
                    )
                    if (i + 1) % 10 == 0:
                        print(f"  {i + 1}/{len(chunks)} 임베딩 완료...")
                itotal += len(chunks)
                print(f"  완료!")
            print(f"\n계면계 완료! 총 {itotal}개 청크 저장됨")

    print(f"\n전체 완료!")


if __name__ == "__main__":
    main()
